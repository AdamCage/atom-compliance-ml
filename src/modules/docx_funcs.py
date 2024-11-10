from pathlib import Path
import re
from typing import Any
from collections import Counter

import docx
from docx import Document
import difflib
from nltk.corpus import stopwords
import pandas as pd


def extract_text_from_docx(docx: docx.document.Document) -> list[str]:
    return [para.text for para in docx.paragraphs]


def extract_elements(lines: list[str], subheader: str) -> str:
    result = []
    start_found = False

    for i, element in enumerate(lines):
        if not start_found and element.lower().strip().replace(" ", "_").startswith(subheader):
            start_found = True
            result.append(element)
            continue

        if start_found:
            if ':' in element:
                break
            result.append(element)

    return "\n".join(result)


def get_differences(uc_text: str, ssts_text: str, diraction: str) -> str:
    stopws = set(stopwords.words('english'))

    diff = list(
        difflib.ndiff(
            [w for w in uc_text.split() if w not in stopws],
            [w for w in ssts_text.split() if w not in stopws]
        )
    )

    diff_stat = Counter([word.replace("+", "") for word in diff if word.startswith(diraction)])

    res = (
        f'Next words are missing in {"UC-text" if diraction == "-" else "SSTS-text"}:\n'
        + "".join(f'{w}: {c}\n' for w, c in diff_stat.items())
    )
    
    return res.replace(",", "")


def create_ds_row(fid: str, hmi_dir: Path, ssts_dir: Path, columns: list[str]) -> dict[str: Any]:
    try:
        hmi_docx = Document(hmi_dir / f'UC-{fid}.docx')

    except Exception as e:
        print(f'Warning! {e}')
        hmi_docx = None

    try:
        ssts_docx = Document(ssts_dir / f'SSTS-{fid}.docx')

    except Exception as e:
        print(f'Warning! {e}')
        ssts_docx = None

    row = dict().fromkeys(columns)

    if hmi_docx is not None:
        uc_lines = extract_text_from_docx(hmi_docx)

        row = dict().fromkeys(columns)

        row["id"] = fid
        row["case_name"] = re.sub(r"\[I-\d+\]", "", re.sub(r'\$\$.*?\$\$|\s*[\xa0]+\s*', ' ', hmi_docx.paragraphs[0].text)).strip()
        row["full_uc_text"] = "\n".join(uc_lines)
        row["full_ssts_text"] = "\n".join(extract_text_from_docx(ssts_docx)) if ssts_docx is not None else None
        row["main_scenario"] = extract_elements(uc_lines, "main_scenario")
        row["goal"] = extract_elements(uc_lines, "goal")
        row["preconditions"] = extract_elements(uc_lines, "preconditions")
        row["postconditions"] = extract_elements(uc_lines, "postconditions")
        row["other"] = (
            extract_elements(uc_lines, "description")
            + extract_elements(uc_lines, "scope")
            + extract_elements(uc_lines, "actors")
            + extract_elements(uc_lines, "requirements")
            + extract_elements(uc_lines, "trigger")
            + extract_elements(uc_lines, "use_case_title")
            + extract_elements(uc_lines, "tigger")
            + extract_elements(uc_lines, "components")
            + extract_elements(uc_lines, "function_logic")
            + extract_elements(uc_lines, "additional_info")
            + extract_elements(uc_lines, "display")
            + extract_elements(uc_lines, "notification")
            + extract_elements(uc_lines, "use_case")
            + extract_elements(uc_lines, "triggers")
            + extract_elements(uc_lines, "requirenments")
        )
        row["alternative_scenario"] = (
            extract_elements(uc_lines, "alternative_scenario_a")
            + extract_elements(uc_lines, "alternative_scenario_b")
            + extract_elements(uc_lines, "alternative_scenario_c")
            + extract_elements(uc_lines, "alternative_scenario")
            + extract_elements(uc_lines, "alternative_scenarios")
            + extract_elements(uc_lines, "alternative_scenario_a_a(turn_off_the_hotspot)")
        )

        row["differences"] = get_differences(
            "\n".join(uc_lines),
            "\n".join(extract_text_from_docx(ssts_docx)) if ssts_docx is not None else "",
            "-"
        )
        row["description"] = get_differences(
            "\n".join(uc_lines),
            "\n".join(extract_text_from_docx(ssts_docx)) if ssts_docx is not None else "",
            "+"
        )
        row["complience_level"] = None

    return row


def create_ds(uc_data_dir: Path, ssts_data_dir: Path, columns: list[str]) -> pd.DataFrame:
    hmi_files = [file.name for file in uc_data_dir.iterdir() if file.is_file()]
    ssts_files = [file.name for file in ssts_data_dir.iterdir() if file.is_file()]

    ids = set(
        [
            "".join(char for char in fname if char.isdigit())
            for fname in (
                hmi_files + ssts_files
        )
        ]
    )

    ds_rows = [
        create_ds_row(fid, uc_data_dir, ssts_data_dir, columns)
        for fid in ids
    ]

    return pd.DataFrame(ds_rows, columns=columns)
